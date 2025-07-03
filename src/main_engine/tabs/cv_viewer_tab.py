"""Tab xem CV trực tiếp mà không cần tải xuống."""

import os
import base64
from pathlib import Path
import streamlit as st
import pandas as pd

from modules.config import ATTACHMENT_DIR, OUTPUT_CSV


def render() -> None:
    """Render UI for viewing CV files directly."""
    st.subheader("🔍 Xem CV trực tiếp")
    
    # Lấy danh sách file CV trong thư mục attachments
    cv_files = [
        f for f in os.listdir(ATTACHMENT_DIR)
        if f.lower().endswith(('.pdf', '.docx'))
    ]
    
    if not cv_files:
        st.info("Không có file CV nào trong thư mục attachments.")
        st.markdown("""
        💡 **Gợi ý:** 
        - Đi đến tab "Lấy & Xử lý CV" để tải CV từ email
        - Hoặc đi đến tab "Single File" để tải lên CV từ máy tính
        """)
        return
    
    # Sắp xếp file theo thời gian sửa đổi (mới nhất trước)
    cv_files.sort(key=lambda f: os.path.getmtime(ATTACHMENT_DIR / f), reverse=True)
    
    # Tạo selectbox với thông tin thêm
    file_options = []
    for f in cv_files:
        file_path = ATTACHMENT_DIR / f
        size_kb = file_path.stat().st_size / 1024
        file_type = "📄 PDF" if f.lower().endswith('.pdf') else "📝 DOCX"
        file_options.append(f"{file_type} {f} ({size_kb:.1f} KB)")
    
    selected_display = st.selectbox(
        "Chọn CV để xem:",
        options=file_options,
        help="Chọn file CV bạn muốn xem trực tiếp trong trình duyệt"
    )
    
    if selected_display:
        # Lấy tên file thực từ display name
        selected_file = selected_display.split(' ', 2)[2].split(' (')[0]
        file_path = ATTACHMENT_DIR / selected_file
        
        # Hiển thị thông tin file
        file_size = file_path.stat().st_size / 1024  # KB
        file_type = "PDF" if selected_file.lower().endswith('.pdf') else "DOCX"
        
        st.markdown(f"""
        <div style="background-color: #f0f2f6; padding: 10px; border-radius: 5px; margin: 10px 0;">
        📁 <strong>File:</strong> {selected_file}<br>
        📊 <strong>Kích thước:</strong> {file_size:.1f} KB<br>
        🗂️ <strong>Loại:</strong> {file_type}
        </div>
        """, unsafe_allow_html=True)
        
        # Hiển thị thông tin CV nếu có trong kết quả
        display_cv_analysis_info(selected_file)
        
        # Tạo tabs cho các chức năng xem khác nhau
        view_tab1, view_tab2, view_tab3 = st.tabs(["📖 Xem trực tiếp", "📄 Nội dung text", "⬇️ Tải xuống"])
        
        with view_tab1:
            if selected_file.lower().endswith('.pdf'):
                st.markdown("### 📖 Xem PDF trực tiếp:")
                display_pdf_viewer(file_path)
            elif selected_file.lower().endswith('.docx'):
                st.markdown("### 📖 Xem DOCX trực tiếp:")
                display_docx_viewer(file_path)
            else:
                st.info("🔍 Tính năng xem trực tiếp chỉ hỗ trợ file PDF và DOCX. Vui lòng sử dụng tab 'Nội dung text' hoặc 'Tải xuống'.")
        
        with view_tab2:
            st.markdown("### 📄 Nội dung text đã trích xuất:")
            display_text_content(file_path)
        
        with view_tab3:
            st.markdown("### ⬇️ Tải file về máy:")
            create_download_button(file_path)


def display_cv_analysis_info(selected_file: str) -> None:
    """Hiển thị thông tin CV đã được phân tích."""
    if os.path.exists(OUTPUT_CSV):
        try:
            df = pd.read_csv(OUTPUT_CSV, encoding="utf-8-sig", keep_default_na=False)
            if "Nguồn" in df.columns:
                cv_info = df[df["Nguồn"] == selected_file]
                if not cv_info.empty:
                    with st.expander("📋 Thông tin đã phân tích", expanded=False):
                        # Tạo 2 cột để hiển thị thông tin
                        col1, col2 = st.columns(2)
                        
                        info_dict = cv_info.iloc[0].to_dict()
                        items = [(k, v) for k, v in info_dict.items() if k != "Nguồn" and str(v).strip() != ""]
                        
                        # Chia đôi danh sách để hiển thị 2 cột
                        mid = len(items) // 2
                        
                        with col1:
                            for k, v in items[:mid]:
                                st.write(f"**{k}:** {v}")
                        
                        with col2:
                            for k, v in items[mid:]:
                                st.write(f"**{k}:** {v}")
                else:
                    st.info("ℹ️ CV này chưa được phân tích. Vui lòng chạy phân tích ở tab 'Lấy & Xử lý CV'.")
        except Exception as e:
            st.warning(f"⚠️ Không thể đọc thông tin CV: {e}")


def display_pdf_viewer(file_path: Path) -> None:
    """Hiển thị PDF viewer nhúng trong Streamlit."""
    try:
        with open(file_path, "rb") as f:
            pdf_data = f.read()
        
        # Encode PDF data to base64
        pdf_base64 = base64.b64encode(pdf_data).decode('utf-8')
        
        # Tạo HTML cho PDF viewer với controls
        pdf_viewer_html = f"""
        <div style="width: 100%; height: 800px; border: 1px solid #ddd; border-radius: 5px; overflow: hidden;">
            <object data="data:application/pdf;base64,{pdf_base64}" type="application/pdf" width="100%" height="100%">
                <iframe src="data:application/pdf;base64,{pdf_base64}" width="100%" height="100%">
                    <p>Trình duyệt của bạn không hỗ trợ hiển thị PDF. 
                    <a href="data:application/pdf;base64,{pdf_base64}" target="_blank">Click vào đây để mở PDF trong tab mới</a>
                    </p>
                </iframe>
            </object>
        </div>
        """
        
        # Hiển thị PDF viewer
        st.markdown(pdf_viewer_html, unsafe_allow_html=True)
        
        # Thêm button để mở trong tab mới
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(
                f'<a href="data:application/pdf;base64,{pdf_base64}" target="_blank" style="text-decoration: none;">'
                '<button style="background-color: #ff4b4b; color: white; padding: 8px 16px; border: none; border-radius: 4px; cursor: pointer;">'
                '🔗 Mở trong tab mới</button></a>',
                unsafe_allow_html=True
            )
        
        with col2:
            if st.button("🔄 Tải lại PDF", help="Tải lại PDF nếu có vấn đề hiển thị"):
                st.experimental_rerun()
        
        # Thông tin hướng dẫn
        st.markdown(
            """
            <div style="background-color: #e1f5fe; padding: 10px; border-radius: 5px; margin-top: 10px;">
            <small>💡 <strong>Hướng dẫn sử dụng:</strong><br>
            • Sử dụng scroll wheel để phóng to/thu nhỏ<br>
            • Click và kéo để di chuyển PDF<br>
            • Nếu PDF không hiển thị, click "Mở trong tab mới" hoặc sử dụng tab "Tải xuống"</small>
            </div>
            """,
            unsafe_allow_html=True
        )
        
    except Exception as e:
        st.error(f"❌ Lỗi khi hiển thị PDF: {e}")
        st.info("🔧 Vui lòng thử sử dụng tab 'Tải xuống' hoặc 'Nội dung text' để xem CV.")


def create_download_button(file_path: Path) -> None:
    """Tạo button download cho file."""
    try:
        with open(file_path, "rb") as f:
            file_data = f.read()
        
        mime_type = (
            "application/pdf" if file_path.suffix.lower() == ".pdf"
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        file_size = len(file_data)
        file_size_mb = file_size / (1024 * 1024)
        
        # Hiển thị thông tin file
        st.markdown(f"""
        <div style="background-color: #e8f5e8; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
        📎 <strong>File sẵn sàng tải xuống:</strong><br>
        🔹 Tên: {file_path.name}<br>
        🔹 Kích thước: {file_size_mb:.2f} MB<br>
        🔹 Định dạng: {file_path.suffix.upper()}
        </div>
        """, unsafe_allow_html=True)
        
        st.download_button(
            label=f"📥 Tải xuống {file_path.name}",
            data=file_data,
            file_name=file_path.name,
            mime=mime_type,
            help="Click để tải file về máy tính",
            use_container_width=True
        )
        
        # Thêm hướng dẫn
        st.markdown("""
        <div style="background-color: #fff3cd; padding: 10px; border-radius: 5px; margin-top: 10px;">
        <small>💡 <strong>Lưu ý:</strong><br>
        • File sẽ được tải về thư mục Downloads mặc định<br>
        • Tùy thuộc vào trình duyệt, bạn có thể được hỏi vị trí lưu file<br>
        • Sau khi tải xong, bạn có thể mở file bằng ứng dụng phù hợp</small>
        </div>
        """, unsafe_allow_html=True)
        
    except Exception as e:
        st.error(f"❌ Lỗi khi tạo link tải xuống: {e}")
        st.info("🔧 Vui lòng thử lại hoặc liên hệ quản trị viên nếu vấn đề tiếp tục xảy ra.")


def display_text_content(file_path: Path) -> None:
    """Hiển thị nội dung text đã trích xuất từ CV."""
    try:
        from modules.cv_processor import CVProcessor
        
        # Tạo spinner để hiển thị quá trình xử lý
        with st.spinner("🔍 Đang trích xuất text từ CV..."):
            processor = CVProcessor()
            text_content = processor.extract_text(str(file_path))
        
        if text_content.strip():
            # Thống kê về nội dung
            word_count = len(text_content.split())
            char_count = len(text_content)
            line_count = len(text_content.split('\n'))
            
            st.markdown(f"""
            <div style="background-color: #f0f2f6; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
            📊 <strong>Thống kê nội dung:</strong> {word_count} từ | {char_count} ký tự | {line_count} dòng
            </div>
            """, unsafe_allow_html=True)
            
            # Hiển thị text trong text area có thể cuộn với chiều cao tự động
            height = min(max(line_count * 20, 200), 600)  # Chiều cao từ 200px đến 600px
            
            st.text_area(
                "Nội dung CV đã trích xuất:",
                value=text_content,
                height=height,
                help="Nội dung text đã được trích xuất từ CV",
                key=f"text_content_{file_path.name}"
            )
            
            # Thêm button copy text (sử dụng JavaScript)
            copy_button_html = """
            <button onclick="copyTextToClipboard()" 
                    style="background-color: #00c851; color: white; padding: 5px 10px; border: none; border-radius: 3px; cursor: pointer;">
            📋 Copy text
            </button>
            <script>
            function copyTextToClipboard() {
                const textArea = document.querySelector('textarea[aria-label="Nội dung CV đã trích xuất:"]');
                if (textArea) {
                    textArea.select();
                    document.execCommand('copy');
                    alert('Đã copy text vào clipboard!');
                }
            }
            </script>
            """
            st.markdown(copy_button_html, unsafe_allow_html=True)
            
        else:
            st.warning("⚠️ Không thể trích xuất text từ file này.")
            st.markdown("""
            **Có thể do:**
            - File bị hỏng hoặc bảo vệ
            - Định dạng file không được hỗ trợ đầy đủ
            - File chỉ chứa hình ảnh (scan) mà không có text
            
            **Gợi ý:** Thử sử dụng tab 'Tải xuống' để mở file bằng ứng dụng khác.
            """)
            
    except Exception as e:
        st.error(f"❌ Lỗi khi trích xuất text: {e}")
        st.info("🔧 Vui lòng thử tải file xuống và mở bằng ứng dụng chuyên dụng.")


def display_docx_viewer(file_path: Path) -> None:
    """Hiển thị nội dung DOCX dưới dạng HTML."""
    try:
        import docx
        from docx.shared import Inches
        
        with st.spinner("📄 Đang xử lý file DOCX..."):
            doc = docx.Document(file_path)
            
            # Tạo HTML content từ DOCX
            html_content = "<div style='background: white; padding: 20px; border: 1px solid #ddd; border-radius: 5px;'>"
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Xử lý định dạng text cơ bản
                    text = paragraph.text
                    style = paragraph.style.name.lower()
                    
                    if 'heading' in style:
                        level = '2'  # Default heading level
                        if 'heading 1' in style:
                            level = '1'
                        elif 'heading 2' in style:
                            level = '2'
                        elif 'heading 3' in style:
                            level = '3'
                        html_content += f"<h{level} style='color: #1f4e79; margin-top: 20px; margin-bottom: 10px;'>{text}</h{level}>"
                    else:
                        html_content += f"<p style='margin-bottom: 10px; line-height: 1.6;'>{text}</p>"
                        
            html_content += "</div>"
            
            # Hiển thị HTML
            st.markdown(html_content, unsafe_allow_html=True)
            
            # Thống kê document
            para_count = len([p for p in doc.paragraphs if p.text.strip()])
            word_count = sum(len(p.text.split()) for p in doc.paragraphs)
            
            st.markdown(f"""
            <div style="background-color: #e1f5fe; padding: 10px; border-radius: 5px; margin-top: 10px;">
            📊 <strong>Thống kê document:</strong> {para_count} đoạn văn | {word_count} từ
            </div>
            """, unsafe_allow_html=True)
            
    except Exception as e:
        st.error(f"❌ Lỗi khi hiển thị DOCX: {e}")
        st.info("🔧 Vui lòng thử sử dụng tab 'Nội dung text' hoặc 'Tải xuống' để xem CV.")
